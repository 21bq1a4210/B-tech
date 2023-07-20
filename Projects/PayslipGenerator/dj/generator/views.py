from django.shortcuts import render
import os
from docx2pdf import convert
from docx import Document
import pandas as pd
from playground import views
from playground.models import Employee_pdf
#from base64 import b85encode

def pdfgenerator():
    def number_to_word(number):
        def get_word(n):
            words={ 0:"", 1:"One", 2:"Two", 3:"Three", 4:"Four", 5:"Five", 6:"Six", 7:"Seven", 8:"Eight", 9:"Nine", 10:"Ten", 11:"Eleven", 12:"Twelve", 13:"Thirteen", 14:"Fourteen", 15:"Fifteen", 16:"Sixteen", 17:"Seventeen", 18:"Eighteen", 19:"Nineteen", 20:"Twenty", 30:"Thirty", 40:"Forty", 50:"Fifty", 60:"Sixty", 70:"Seventy", 80:"Eighty", 90:"Ninty" }
            if n<=20:
                return words[n]
            else:
                ones=n%10
                tens=n-ones
                return words[tens]+" "+words[ones]
                
        def get_all_word(n):
            d=[100,10,100,100]
            v=["","Hundred And","Thousand","lakh"]
            w=[]
            for i,x in zip(d,v):
                t=get_word(n%i)
                if t!="":
                    t+=" "+x
                w.append(t.rstrip(" "))
                n=n//i
            w.reverse()
            w=' '.join(w).strip()
            if w.endswith("And"):
                w=w[:-3]
            return w
        arr=str(number).split(".")
        number=int(arr[0])
        crore=number//10000000
        number=number%10000000
        word=""
        if crore>0:
            word+=get_all_word(crore)
            word+=" crore "
        word+=get_all_word(number).strip()+" Rupees"
        if len(arr)>1:
            if len(arr[1])==1:
                arr[1]+="0"
            word+=" and "+get_all_word(int(arr[1]))+" paisa"
        return word
    data=views.exceldata()
    nrow=len(data.index)
    slip=Document("PAY SLIP.docx")
    maindir=os.getcwd()
    cwd=maindir
    folders=os.listdir(cwd)
    if "PDF's Folder" not in folders:
        pdfdir=os.path.join(cwd,"PDF's Folder")
        os.mkdir(pdfdir)
    else:
        pdfdir=os.path.join(cwd,"PDF's Folder")
    folders =os.listdir(pdfdir)
    list0=[[2,1],[1,1],[1,3],[2,3],[3,1],[3,3],[4,1],[4,3],[5,1],[5,3]]
    list1=[[1,1],[2,1],[3,1],[4,1],[5,1],[10,1],[1,3],[2,3],[3,3],[4,3],[5,3],[6,3],[7,3],[8,3],[9,3],[10,3]]
    for i in range(1,nrow):
        cwd=pdfdir
        l0=0
        l1=0
        for j in range(10):
            x,y=list0[j]
            if j in [1,4,6,7,8]:
                slip.tables[0].cell(x,y).text=str((data.iat[i,j]))
            else:
                slip.tables[0].cell(x,y).text=str(int(data.iat[i,j]))
        for k in range(10,26):
            x,y=list1[k-10]
            slip.tables[1].cell(x,y).text=str((data.iat[i,k]))
        words=number_to_word(int(data.iat[i,k]))
        slip.tables[1].cell(11,1).text=words
        filename=str(data.iat[i,7]).upper()+"-"+str(int(data.iat[i,9])).upper()
        emp_id=str(int(data.iat[i,0]))
        if emp_id not in folders:
            cwd=os.path.join(cwd,emp_id)
            os.mkdir(cwd)
            pdf=Employee_pdf(employee_id=emp_id,employee_payslip_dir=cwd)
            pdf.save()
        else:
            cwd=os.path.join(cwd,emp_id)
        curr=cwd+"\\"+filename+".docx"
        slip.save(curr)
        dest=cwd+"\\"+filename+".pdf"
        convert(curr,dest)
        os.remove(curr)
    return